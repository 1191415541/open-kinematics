"""Unified Word report export helpers for GUI workbenches."""

from __future__ import annotations

from dataclasses import dataclass, replace
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Literal

from matplotlib.figure import Figure

from kinematics.gui.steering.plotting import (
    WHEEL_RADIUS,
    WHEEL_WIDTH,
    draw_curve_plot,
    draw_steering_preview,
    draw_three_segment_steering_preview,
)
from kinematics.gui.suspension.plotting import (
    SuspensionPreviewRenderer,
    draw_suspension_curve_plot,
    draw_suspension_preview,
)
from kinematics.gui.suspension.reporting import (
    build_metric_range_rows as build_suspension_metric_range_rows,
)
from kinematics.gui.suspension.reporting import (
    summarize_suspension_curve,
)
from kinematics.gui.suspension.workbench import (
    SuspensionProject,
    SuspensionSweepResult,
    solve_suspension_project,
)
from kinematics.gui.suspension.workbench import (
    curve_specs_for_plot as suspension_curve_specs_for_plot,
)
from kinematics.steering.two_segment import solve_two_segment_steering
from kinematics.steering.workbench import (
    SteeringProject,
    hardpoints_from_rows,
    solve_steering_project,
    solve_three_segment_project,
    sweep_steering_project,
    three_segment_geometry_from_rows,
)
from kinematics.steering.workbench import (
    curve_specs_for_plot as steering_curve_specs_for_plot,
)

ReportScope = Literal["suspension", "steering", "combined"]
ReportImageKey = Literal[
    "suspension_preview",
    "steering_preview",
]

STEERING_OUTPUT_LABELS = {
    "input_value": "Input value [deg]",
    "pitman_angle_deg": "Pitman angle [deg]",
    "left_bellcrank_angle_deg": "Left bellcrank angle [deg]",
    "right_bellcrank_angle_deg": "Right bellcrank angle [deg]",
    "left_wheel_angle_deg": "Left wheel angle [deg]",
    "right_wheel_angle_deg": "Right wheel angle [deg]",
    "left_minus_right_deg": "Left minus right [deg]",
    "ackermann_rate_pct": "Ackermann rate [%]",
    "max_left_turn_left_wheel_angle_deg": "Max left turn left wheel [deg]",
    "max_left_turn_right_wheel_angle_deg": "Max left turn right wheel [deg]",
    "max_right_turn_left_wheel_angle_deg": "Max right turn left wheel [deg]",
    "max_right_turn_right_wheel_angle_deg": "Max right turn right wheel [deg]",
    "left_tie_rod_residual": "Left tie rod residual",
    "right_tie_rod_residual": "Right tie rod residual",
}
STEERING_SUMMARY_METRICS = (
    "left_wheel_angle_deg",
    "right_wheel_angle_deg",
    "left_minus_right_deg",
    "ackermann_rate_pct",
    "left_tie_rod_residual",
    "right_tie_rod_residual",
)


@dataclass(frozen=True)
class ReportCurveSelection:
    """One user-selected report curve combination."""

    x_output: str
    y_output: str
    label: str = ""


@dataclass(frozen=True)
class ReportExportOptions:
    """User-selected report export options."""

    scope: ReportScope
    include_images: tuple[ReportImageKey, ...]
    suspension_curves: tuple[ReportCurveSelection, ...] | None = None
    steering_curves: tuple[ReportCurveSelection, ...] | None = None


@dataclass(frozen=True)
class SteeringCurveReportSummary:
    """Summary information for one exported steering curve."""

    label: str
    x_output: str
    y_output: str
    sample_count: int
    x_start: float
    x_end: float
    y_min: float
    y_max: float
    y_min_at_x: float
    y_max_at_x: float
    trend: str
    has_turning_point: bool
    crosses_zero: bool

    def description(self) -> str:
        x_label = steering_output_label(self.x_output)
        y_label = steering_output_label(self.y_output)
        trend_text = {
            "increasing": "is overall increasing",
            "decreasing": "is overall decreasing",
            "flat": "stays nearly flat",
            "mixed": "is non-monotonic",
        }[self.trend]
        turning_text = (
            " and includes at least one turning point" if self.has_turning_point else ""
        )
        zero_text = (
            f" {y_label} crosses zero during the sweep."
            if self.crosses_zero
            else f" {y_label} does not cross zero during the sweep."
        )
        return (
            f"{self.label} uses {self.sample_count} solved steps from "
            f"{_format_number(self.x_start)} to {_format_number(self.x_end)} "
            f"{x_label}. {y_label} ranges from {_format_number(self.y_min)} to "
            f"{_format_number(self.y_max)}; the minimum occurs near "
            f"{x_label}={_format_number(self.y_min_at_x)} and the maximum near "
            f"{x_label}={_format_number(self.y_max_at_x)}. The curve {trend_text}"
            f"{turning_text}.{zero_text}"
        )


def export_gui_report_docx(
    path: str | Path,
    *,
    options: ReportExportOptions,
    suspension_project: SuspensionProject | None = None,
    steering_project: SteeringProject | None = None,
    suspension_source_path: str | Path | None = None,
    steering_source_path: str | Path | None = None,
    generated_at: datetime | None = None,
) -> None:
    """Export one GUI report to Word with chapter structure and TOC."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ModuleNotFoundError as exc:
        raise ModuleNotFoundError(
            "Report export requires python-docx. "
            "Install/run with: uv run --extra viz kinematics gui"
        ) from exc

    if options.scope in {"suspension", "combined"} and suspension_project is None:
        raise ValueError("Suspension report export requires a suspension project")
    if options.scope in {"steering", "combined"} and steering_project is None:
        raise ValueError("Steering report export requires a steering project")

    document = Document()
    generated = generated_at or datetime.now().astimezone()
    document.add_heading("Kinematics GUI Report", level=0)
    document.add_paragraph(
        "Generated "
        f"{generated:%Y-%m-%d %H:%M %Z}. This report is organized by chapter and "
        "summarizes the selected GUI workbench data, figures, and kinematics trends."
    )

    document.add_heading("Table Of Contents", level=1)
    _add_table_of_contents(document)
    document.add_page_break()

    summary_rows: list[tuple[str, str, float | None, float | None, str]] = []

    if options.scope in {"suspension", "combined"}:
        assert suspension_project is not None
        sweep = solve_suspension_project(suspension_project)
        suspension_curves = _resolve_suspension_report_curves(
            project=suspension_project,
            selections=options.suspension_curves,
        )
        _append_suspension_section(
            document,
            project=suspension_project,
            sweep=sweep,
            curves=suspension_curves,
            source_path=suspension_source_path,
            include_preview="suspension_preview" in options.include_images,
            summary_rows=summary_rows,
            image_width=Inches(6.4),
        )
        if options.scope == "combined":
            document.add_page_break()

    if options.scope in {"steering", "combined"}:
        assert steering_project is not None
        steering_rows = sweep_steering_project(steering_project, skip_unreachable=True)
        steering_curves = _resolve_steering_report_curves(
            project=steering_project,
            selections=options.steering_curves,
        )
        _append_steering_section(
            document,
            project=steering_project,
            rows=steering_rows,
            curves=steering_curves,
            source_path=steering_source_path,
            include_preview="steering_preview" in options.include_images,
            summary_rows=summary_rows,
            image_width=Inches(6.4),
        )

    document.add_page_break()
    document.add_heading("Kinematic Parameter Summary", level=1)
    document.add_paragraph(
        "The table below summarizes the overall parameter variation captured in the "
        "selected report scope."
    )
    summary_table = document.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"
    header = summary_table.rows[0].cells
    header[0].text = "Domain"
    header[1].text = "Parameter"
    header[2].text = "Minimum"
    header[3].text = "Maximum"
    header[4].text = "Comment"
    for domain, label, minimum, maximum, comment in summary_rows:
        row = summary_table.add_row().cells
        row[0].text = domain
        row[1].text = label
        row[2].text = _format_optional_number(minimum)
        row[3].text = _format_optional_number(maximum)
        row[4].text = comment

    document.save(Path(path))


def render_suspension_preview_png(
    project: SuspensionProject,
    *,
    width: float = 6.5,
    height: float = 5.0,
    dpi: int = 160,
) -> bytes:
    """Render the suspension design-condition preview to PNG bytes."""
    figure = Figure(figsize=(width, height), dpi=dpi)
    axis = figure.add_subplot(111, projection="3d")
    renderer = SuspensionPreviewRenderer()
    suspension = project.build_suspension()
    draw_suspension_preview(
        axis,
        suspension,
        suspension.initial_state(),
        preserve_view=False,
        renderer=renderer,
        preview_mode=False,
    )
    return _save_figure_to_png(figure)


def render_suspension_curves_png(
    rows: list[dict[str, float | bool | None]],
    curves: list[tuple[str, str, str]],
    *,
    width: float = 7.0,
    height: float = 4.2,
    dpi: int = 160,
) -> bytes:
    """Render suspension output curves to PNG bytes."""
    figure = Figure(figsize=(width, height), dpi=dpi)
    axis = figure.add_subplot(111)
    draw_suspension_curve_plot(axis, rows, curves)
    return _save_figure_to_png(figure)


def render_steering_preview_png(
    project: SteeringProject,
    *,
    width: float = 7.0,
    height: float = 5.5,
    dpi: int = 160,
) -> bytes:
    """Render the steering design/current preview to PNG bytes."""
    figure = Figure(figsize=(width, height), dpi=dpi)
    axis = figure.add_subplot(111)
    if project.linkage_type == "three_segment":
        geometry = three_segment_geometry_from_rows(project.hardpoints)
        design_state = solve_three_segment_project(
            replace(
                project,
                input_mode="left_bellcrank_angle",
                input_value=0.0,
            )
        )
        current_state = solve_three_segment_project(project)
        draw_three_segment_steering_preview(
            axis,
            geometry,
            design_state,
            current_state,
            preserve_view=False,
            wheel_radius=project.wheel_radius,
            wheel_width=project.wheel_width,
        )
    else:
        hardpoints = hardpoints_from_rows(project.hardpoints)
        design_state = solve_two_segment_steering(hardpoints, 0.0)
        current_state, _outputs = solve_steering_project(project, include_limits=False)
        draw_steering_preview(
            axis,
            hardpoints,
            design_state,
            current_state,
            preserve_view=False,
            wheel_radius=project.wheel_radius or WHEEL_RADIUS,
            wheel_width=project.wheel_width or WHEEL_WIDTH,
        )
    return _save_figure_to_png(figure)


def render_steering_curves_png(
    rows: list[dict[str, float]],
    curves: list[tuple[str, str, str]],
    *,
    width: float = 7.0,
    height: float = 4.2,
    dpi: int = 160,
) -> bytes:
    """Render steering output curves to PNG bytes."""
    figure = Figure(figsize=(width, height), dpi=dpi)
    axis = figure.add_subplot(111)
    draw_curve_plot(axis, rows, curves)
    return _save_figure_to_png(figure)


def summarize_steering_curve(
    rows: list[dict[str, float]],
    *,
    x_output: str,
    y_output: str,
    label: str,
) -> SteeringCurveReportSummary | None:
    """Summarize one steering output curve for report prose."""
    samples = [
        (float(row[x_output]), float(row[y_output]))
        for row in rows
        if x_output in row and y_output in row
    ]
    if not samples:
        return None
    x_values = [x_value for x_value, _y_value in samples]
    y_values = [y_value for _x_value, y_value in samples]
    y_min = min(y_values)
    y_max = max(y_values)
    y_min_index = y_values.index(y_min)
    y_max_index = y_values.index(y_max)
    trend, has_turning_point, crosses_zero = _curve_characteristics(y_values)
    curve_label = (
        label.strip()
        or f"{steering_output_label(y_output)} vs {steering_output_label(x_output)}"
    )
    return SteeringCurveReportSummary(
        label=curve_label,
        x_output=x_output,
        y_output=y_output,
        sample_count=len(samples),
        x_start=x_values[0],
        x_end=x_values[-1],
        y_min=y_min,
        y_max=y_max,
        y_min_at_x=x_values[y_min_index],
        y_max_at_x=x_values[y_max_index],
        trend=trend,
        has_turning_point=has_turning_point,
        crosses_zero=crosses_zero,
    )


def build_steering_metric_range_rows(
    rows: list[dict[str, float]],
) -> list[tuple[str, float, float]]:
    """Return min/max rows for key steering metrics."""
    range_rows: list[tuple[str, float, float]] = []
    for metric_name in STEERING_SUMMARY_METRICS:
        values = [float(row[metric_name]) for row in rows if metric_name in row]
        if not values:
            continue
        range_rows.append(
            (steering_output_label(metric_name), min(values), max(values))
        )
    return range_rows


def steering_output_label(output_name: str) -> str:
    """Return a display label for one steering output key."""
    known_label = STEERING_OUTPUT_LABELS.get(output_name)
    if known_label is not None:
        return known_label
    if output_name.endswith("_deg"):
        return output_name.removesuffix("_deg").replace("_", " ").title() + " [deg]"
    if output_name.endswith("_pct"):
        return output_name.removesuffix("_pct").replace("_", " ").title() + " [%]"
    return output_name.replace("_", " ").title()


def _append_suspension_section(
    document,
    *,
    project: SuspensionProject,
    sweep: SuspensionSweepResult,
    curves: list[tuple[str, str, str]],
    source_path: str | Path | None,
    include_preview: bool,
    summary_rows: list[tuple[str, str, float | None, float | None, str]],
    image_width,
) -> None:
    document.add_heading("Suspension", level=1)
    document.add_heading("Project Summary", level=2)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Item"
    summary_table.rows[0].cells[1].text = "Value"
    source_text = str(source_path) if source_path is not None else "Unsaved GUI project"
    for item, value in (
        ("Project", project.name),
        ("Suspension type", project.suspension_type.replace("_", " ").title()),
        ("Source file", source_text),
        ("Sweep steps", str(project.settings.steps)),
        (
            "Sweep travel [mm]",
            f"{_format_number(project.settings.start)} to "
            f"{_format_number(project.settings.stop)}",
        ),
    ):
        row = summary_table.add_row().cells
        row[0].text = item
        row[1].text = value

    document.add_heading("Analysis Summary", level=2)
    document.add_paragraph(
        f"The suspension sweep contains {len(sweep.rows)} solved states. "
        "This chapter summarizes the current travel sweep, selected figures, "
        "and the main kinematic parameter trends."
    )

    if include_preview:
        document.add_heading("Design Preview", level=2)
        document.add_paragraph(
            "The following image shows the suspension geometry at design condition."
        )
        document.add_picture(
            BytesIO(render_suspension_preview_png(project)),
            width=image_width,
        )

    document.add_heading("Curve Results", level=2)
    if not curves:
        document.add_paragraph("No suspension curves were selected for export.")
    for index, (x_output, y_output, label) in enumerate(curves, start=1):
        summary = summarize_suspension_curve(
            sweep.rows,
            x_output=x_output,
            y_output=y_output,
            label=label,
        )
        if summary is None:
            continue
        document.add_heading(f"Curve {index}: {summary.label}", level=3)
        document.add_paragraph(
            "The following figure shows the selected suspension curve only."
        )
        document.add_picture(
            BytesIO(
                render_suspension_curves_png(
                    sweep.rows,
                    [(x_output, y_output, label)],
                )
            ),
            width=image_width,
        )
        document.add_paragraph(summary.description())

    document.add_heading("Parameter Variation Summary", level=2)
    metric_ranges = build_suspension_metric_range_rows(sweep.rows)
    if metric_ranges:
        metric_table = document.add_table(rows=1, cols=3)
        metric_table.style = "Table Grid"
        metric_table.rows[0].cells[0].text = "Parameter"
        metric_table.rows[0].cells[1].text = "Minimum"
        metric_table.rows[0].cells[2].text = "Maximum"
        for label, minimum, maximum in metric_ranges:
            row = metric_table.add_row().cells
            row[0].text = label
            row[1].text = _format_number(minimum)
            row[2].text = _format_number(maximum)
            summary_rows.append(
                (
                    "Suspension",
                    label,
                    minimum,
                    maximum,
                    _range_comment(minimum, maximum),
                )
            )
    else:
        document.add_paragraph("No numeric suspension metrics are available.")


def _append_steering_section(
    document,
    *,
    project: SteeringProject,
    rows: list[dict[str, float]],
    curves: list[tuple[str, str, str]],
    source_path: str | Path | None,
    include_preview: bool,
    summary_rows: list[tuple[str, str, float | None, float | None, str]],
    image_width,
) -> None:
    document.add_heading("Steering", level=1)
    document.add_heading("Project Summary", level=2)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Item"
    summary_table.rows[0].cells[1].text = "Value"
    source_text = str(source_path) if source_path is not None else "Unsaved GUI project"
    for item, value in (
        ("Project", project.name),
        ("Linkage type", project.linkage_type.replace("_", " ").title()),
        ("Source file", source_text),
        ("Input mode", project.input_mode),
        (
            "Sweep input range",
            f"{_format_number(project.sweep_min)} to "
            f"{_format_number(project.sweep_max)}",
        ),
        ("Sweep step", _format_number(project.sweep_step)),
        ("Wheelbase [mm]", _format_number(project.wheelbase)),
    ):
        row = summary_table.add_row().cells
        row[0].text = item
        row[1].text = value

    document.add_heading("Analysis Summary", level=2)
    document.add_paragraph(
        f"The steering sweep contains {len(rows)} solved states. "
        "This chapter summarizes current steering geometry, output curves, "
        "and steering parameter trends such as wheel-angle split and Ackermann rate."
    )

    if include_preview:
        document.add_heading("Geometry Preview", level=2)
        document.add_paragraph(
            "The following image shows the steering linkage at design condition "
            "and the current solved state."
        )
        document.add_picture(
            BytesIO(render_steering_preview_png(project)),
            width=image_width,
        )

    document.add_heading("Curve Results", level=2)
    if not curves:
        document.add_paragraph("No steering curves were selected for export.")
    for index, (x_output, y_output, label) in enumerate(curves, start=1):
        summary = summarize_steering_curve(
            rows,
            x_output=x_output,
            y_output=y_output,
            label=label,
        )
        if summary is None:
            continue
        document.add_heading(f"Curve {index}: {summary.label}", level=3)
        document.add_paragraph(
            "The following figure shows the selected steering curve only."
        )
        document.add_picture(
            BytesIO(
                render_steering_curves_png(
                    rows,
                    [(x_output, y_output, label)],
                )
            ),
            width=image_width,
        )
        document.add_paragraph(summary.description())

    document.add_heading("Parameter Variation Summary", level=2)
    metric_ranges = build_steering_metric_range_rows(rows)
    if metric_ranges:
        metric_table = document.add_table(rows=1, cols=3)
        metric_table.style = "Table Grid"
        metric_table.rows[0].cells[0].text = "Parameter"
        metric_table.rows[0].cells[1].text = "Minimum"
        metric_table.rows[0].cells[2].text = "Maximum"
        for label, minimum, maximum in metric_ranges:
            row = metric_table.add_row().cells
            row[0].text = label
            row[1].text = _format_number(minimum)
            row[2].text = _format_number(maximum)
            summary_rows.append(
                ("Steering", label, minimum, maximum, _range_comment(minimum, maximum))
            )
    else:
        document.add_paragraph("No numeric steering metrics are available.")


def _add_table_of_contents(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    separate_text = OxmlElement("w:t")
    separate_text.text = (
        "Right-click and update field in Word to refresh the table of contents."
    )
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(separate_text)
    run._r.append(fld_char_end)


def _curve_characteristics(
    y_values: list[float],
) -> tuple[str, bool, bool]:
    span = max(y_values) - min(y_values)
    tolerance = max(abs(span) * 1e-6, 1e-9)
    deltas = [right - left for left, right in zip(y_values, y_values[1:])]
    increasing = all(delta >= -tolerance for delta in deltas)
    decreasing = all(delta <= tolerance for delta in deltas)
    if span <= tolerance:
        trend = "flat"
    elif increasing and not decreasing:
        trend = "increasing"
    elif decreasing and not increasing:
        trend = "decreasing"
    else:
        trend = "mixed"
    directions = [
        1 if delta > tolerance else -1 if delta < -tolerance else 0 for delta in deltas
    ]
    filtered_directions = [direction for direction in directions if direction != 0]
    has_turning_point = any(
        left != right
        for left, right in zip(filtered_directions, filtered_directions[1:])
    )
    crosses_zero = any(abs(value) <= tolerance for value in y_values) or any(
        (left < -tolerance and right > tolerance)
        or (left > tolerance and right < -tolerance)
        for left, right in zip(y_values, y_values[1:])
    )
    return trend, has_turning_point, crosses_zero


def _range_comment(minimum: float, maximum: float) -> str:
    span = maximum - minimum
    if abs(span) <= 1e-9:
        return "Nearly constant across the solved range."
    if span > 0.0:
        return f"Net variation {_format_number(span)}."
    return f"Net variation {_format_number(span)}."


def _resolve_suspension_report_curves(
    *,
    project: SuspensionProject,
    selections: tuple[ReportCurveSelection, ...] | None,
) -> list[tuple[str, str, str]]:
    if selections is None:
        return suspension_curve_specs_for_plot(
            project.curves,
            "wheel_travel_mm",
            "camber_deg",
            "",
        )
    return _report_curve_specs(selections)


def _resolve_steering_report_curves(
    *,
    project: SteeringProject,
    selections: tuple[ReportCurveSelection, ...] | None,
) -> list[tuple[str, str, str]]:
    if selections is None:
        return steering_curve_specs_for_plot(
            project.curves,
            "input_value",
            "left_wheel_angle_deg",
            "",
        )
    return _report_curve_specs(selections)


def _report_curve_specs(
    selections: tuple[ReportCurveSelection, ...],
) -> list[tuple[str, str, str]]:
    return [
        (selection.x_output, selection.y_output, selection.label)
        for selection in selections
    ]


def _save_figure_to_png(figure: Figure) -> bytes:
    buffer = BytesIO()
    figure.savefig(buffer, format="png", dpi=figure.dpi, bbox_inches="tight")
    figure.clear()
    return buffer.getvalue()


def _format_number(value: float) -> str:
    return f"{float(value):.6g}"


def _format_optional_number(value: float | None) -> str:
    if value is None:
        return "-"
    return _format_number(value)
