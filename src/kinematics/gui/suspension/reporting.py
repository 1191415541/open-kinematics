"""Word report export helpers for the suspension GUI."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path

from matplotlib.figure import Figure

from kinematics.gui.suspension.plotting import draw_suspension_curve_plot
from kinematics.gui.suspension.workbench import SuspensionProject, SuspensionSweepResult

OUTPUT_LABELS = {
    "wheel_travel_mm": "Wheel travel [mm]",
    "camber_deg": "Camber [deg]",
    "caster_deg": "Caster [deg]",
    "roadwheel_angle_deg": "Road wheel angle [deg]",
    "toe_deg": "Toe [deg]",
    "kpi_deg": "KPI [deg]",
    "scrub_radius_mm": "Scrub radius [mm]",
    "mechanical_trail_mm": "Mechanical trail [mm]",
    "roll_center_height_mm": "Roll center height [mm]",
    "roll_center_lateral_offset_mm": "Roll center lateral offset [mm]",
    "anti_pitch_pct": "Anti-pitch [%]",
    "track_change_mm": "Track change [mm]",
    "solver_max_residual": "Solver max residual",
}
KEY_REPORT_METRICS = (
    "camber_deg",
    "toe_deg",
    "caster_deg",
    "kpi_deg",
    "scrub_radius_mm",
    "mechanical_trail_mm",
    "roll_center_height_mm",
    "roll_center_lateral_offset_mm",
    "anti_pitch_pct",
    "track_change_mm",
    "solver_max_residual",
)


@dataclass(frozen=True)
class SuspensionCurveReportSummary:
    """Summary information for one exported suspension curve."""

    label: str
    x_output: str
    y_output: str
    sample_count: int
    x_start: float
    x_end: float
    y_start: float
    y_end: float
    x_min: float
    x_max: float
    y_min: float
    y_max: float
    y_min_at_x: float
    y_max_at_x: float
    trend: str
    has_turning_point: bool
    crosses_zero: bool

    def description(self) -> str:
        """Return one narrative description for the curve."""
        x_label = output_label(self.x_output)
        y_label = output_label(self.y_output)
        trend_text = {
            "increasing": "is overall increasing",
            "decreasing": "is overall decreasing",
            "flat": "stays nearly flat",
            "mixed": "is non-monotonic",
        }[self.trend]
        turning_text = (
            " and includes at least one turning point" if self.has_turning_point else ""
        )
        zero_text = (
            f" {y_label} crosses zero during the sweep."
            if self.crosses_zero
            else f" {y_label} does not cross zero during the sweep."
        )
        return (
            f"{self.label} uses {self.sample_count} solved steps from "
            f"{_format_number(self.x_start)} to {_format_number(self.x_end)} "
            f"{x_label}. {y_label} ranges from {_format_number(self.y_min)} to "
            f"{_format_number(self.y_max)}; the minimum occurs near "
            f"{x_label}={_format_number(self.y_min_at_x)} and the maximum near "
            f"{x_label}={_format_number(self.y_max_at_x)}. The curve {trend_text}"
            f"{turning_text}.{zero_text}"
        )


def export_suspension_report_docx(
    path: str | Path,
    *,
    project: SuspensionProject,
    sweep: SuspensionSweepResult,
    curves: list[tuple[str, str, str]],
    source_path: str | Path | None = None,
    generated_at: datetime | None = None,
) -> None:
    """Export a Word report for the current suspension sweep."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ModuleNotFoundError as exc:
        raise ModuleNotFoundError(
            "Report export requires python-docx. "
            "Install/run with: uv run --extra viz kinematics gui"
        ) from exc

    if not sweep.rows:
        raise ValueError("No solved sweep rows are available for report export")

    curve_summaries = [
        summary
        for x_output, y_output, label in curves
        if (
            summary := summarize_suspension_curve(
                sweep.rows,
                x_output=x_output,
                y_output=y_output,
                label=label,
            )
        )
        is not None
    ]
    if not curve_summaries:
        raise ValueError("No numeric curve data is available for report export")

    document = Document()
    generated = generated_at or datetime.now().astimezone()
    source_text = str(source_path) if source_path is not None else "Unsaved GUI project"
    converged_steps = sum(bool(row.get("solver_converged")) for row in sweep.rows)
    peak_residual = max(
        (float(value) for value in _numeric_series(sweep.rows, "solver_max_residual")),
        default=0.0,
    )

    document.add_heading("Suspension Kinematics Report", level=0)
    document.add_paragraph(
        "Generated "
        f"{generated:%Y-%m-%d %H:%M %Z}. This report summarizes the current "
        "suspension project, the solved wheel-travel sweep, and the configured "
        "kinematics curves."
    )

    document.add_heading("Project Summary", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Item"
    summary_table.rows[0].cells[1].text = "Value"
    for item, value in (
        ("Project", project.name),
        ("Suspension type", project.suspension_type.replace("_", " ").title()),
        ("Source file", source_text),
        ("Units", str(project.units.name).title()),
        ("Wheelbase [mm]", _format_number(project.config.wheelbase)),
        (
            "Sweep travel [mm]",
            f"{_format_number(project.settings.start)} to "
            f"{_format_number(project.settings.stop)}",
        ),
        ("Sweep steps", str(project.settings.steps)),
        ("Configured curves", str(len(curve_summaries))),
        ("Solver convergence", f"{converged_steps}/{len(sweep.rows)}"),
    ):
        row = summary_table.add_row().cells
        row[0].text = item
        row[1].text = value

    document.add_heading("Sweep Summary", level=1)
    convergence_note = (
        "All solved steps converged cleanly."
        if converged_steps == len(sweep.rows)
        else "Some steps did not converge; review the numeric outputs before release."
    )
    document.add_paragraph(
        f"The sweep covers {len(sweep.rows)} solved states with a peak solver "
        f"residual of {_format_number(peak_residual)}. {convergence_note}"
    )

    metric_ranges = build_metric_range_rows(sweep.rows)
    if metric_ranges:
        metric_table = document.add_table(rows=1, cols=3)
        metric_table.style = "Table Grid"
        metric_table.rows[0].cells[0].text = "Output"
        metric_table.rows[0].cells[1].text = "Minimum"
        metric_table.rows[0].cells[2].text = "Maximum"
        for label, min_value, max_value in metric_ranges:
            row = metric_table.add_row().cells
            row[0].text = label
            row[1].text = _format_number(min_value)
            row[2].text = _format_number(max_value)

    document.add_heading("Curve Plot", level=1)
    document.add_paragraph(
        "The figure below overlays the currently configured suspension output "
        "curves across the full wheel-travel sweep."
    )
    document.add_picture(
        BytesIO(render_suspension_curve_plot_png(sweep.rows, curves)),
        width=Inches(6.5),
    )

    document.add_heading("Curve Descriptions", level=1)
    for index, summary in enumerate(curve_summaries, start=1):
        document.add_paragraph(f"Curve {index}. {summary.description()}")

    document.save(Path(path))


def render_suspension_curve_plot_png(
    rows: list[dict[str, float | bool | None]],
    curves: list[tuple[str, str, str]],
) -> bytes:
    """Render the configured curve plot to a PNG byte string."""
    figure = Figure(figsize=(7.0, 4.2), dpi=160)
    axis = figure.add_subplot(111)
    draw_suspension_curve_plot(axis, rows, curves)
    buffer = BytesIO()
    figure.savefig(buffer, format="png", dpi=160, bbox_inches="tight")
    figure.clear()
    return buffer.getvalue()


def summarize_suspension_curve(
    rows: list[dict[str, float | bool | None]],
    *,
    x_output: str,
    y_output: str,
    label: str,
) -> SuspensionCurveReportSummary | None:
    """Summarize one suspension output curve for report prose."""
    samples = [
        (float(row[x_output]), float(row[y_output]))
        for row in rows
        if _is_number(row.get(x_output)) and _is_number(row.get(y_output))
    ]
    if not samples:
        return None

    x_values = [x_value for x_value, _y_value in samples]
    y_values = [y_value for _x_value, y_value in samples]
    y_min = min(y_values)
    y_max = max(y_values)
    y_min_index = y_values.index(y_min)
    y_max_index = y_values.index(y_max)
    trend, has_turning_point, crosses_zero = _curve_characteristics(y_values)
    curve_label = (
        label.strip() or f"{output_label(y_output)} vs {output_label(x_output)}"
    )
    return SuspensionCurveReportSummary(
        label=curve_label,
        x_output=x_output,
        y_output=y_output,
        sample_count=len(samples),
        x_start=x_values[0],
        x_end=x_values[-1],
        y_start=y_values[0],
        y_end=y_values[-1],
        x_min=min(x_values),
        x_max=max(x_values),
        y_min=y_min,
        y_max=y_max,
        y_min_at_x=x_values[y_min_index],
        y_max_at_x=x_values[y_max_index],
        trend=trend,
        has_turning_point=has_turning_point,
        crosses_zero=crosses_zero,
    )


def build_metric_range_rows(
    rows: list[dict[str, float | bool | None]],
) -> list[tuple[str, float, float]]:
    """Return min/max rows for the key metrics shown in the report."""
    range_rows: list[tuple[str, float, float]] = []
    for metric_name in KEY_REPORT_METRICS:
        values = _numeric_series(rows, metric_name)
        if not values:
            continue
        range_rows.append((output_label(metric_name), min(values), max(values)))
    return range_rows


def output_label(output_name: str) -> str:
    """Return a display label for one output key."""
    known_label = OUTPUT_LABELS.get(output_name)
    if known_label is not None:
        return known_label
    if output_name.endswith("_deg"):
        return output_name.removesuffix("_deg").replace("_", " ").title() + " [deg]"
    if output_name.endswith("_mm"):
        return output_name.removesuffix("_mm").replace("_", " ").title() + " [mm]"
    return output_name.replace("_", " ").title()


def _curve_characteristics(
    y_values: list[float],
) -> tuple[str, bool, bool]:
    span = max(y_values) - min(y_values)
    tolerance = max(abs(span) * 1e-6, 1e-9)
    deltas = [right - left for left, right in zip(y_values, y_values[1:])]
    increasing = all(delta >= -tolerance for delta in deltas)
    decreasing = all(delta <= tolerance for delta in deltas)
    if span <= tolerance:
        trend = "flat"
    elif increasing and not decreasing:
        trend = "increasing"
    elif decreasing and not increasing:
        trend = "decreasing"
    else:
        trend = "mixed"

    directions = [
        1 if delta > tolerance else -1 if delta < -tolerance else 0 for delta in deltas
    ]
    filtered_directions = [direction for direction in directions if direction != 0]
    has_turning_point = any(
        left != right
        for left, right in zip(filtered_directions, filtered_directions[1:])
    )
    crosses_zero = any(abs(value) <= tolerance for value in y_values) or any(
        (left < -tolerance and right > tolerance)
        or (left > tolerance and right < -tolerance)
        for left, right in zip(y_values, y_values[1:])
    )
    return trend, has_turning_point, crosses_zero


def _numeric_series(
    rows: list[dict[str, float | bool | None]],
    key: str,
) -> list[float]:
    return [float(row[key]) for row in rows if _is_number(row.get(key))]


def _is_number(value: object) -> bool:
    return isinstance(value, int | float) and not isinstance(value, bool)


def _format_number(value: float) -> str:
    return f"{float(value):.6g}"
